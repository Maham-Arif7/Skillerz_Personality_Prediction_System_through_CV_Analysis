"""
extractor.py
------------------------------------------------------------------
Extracts raw text from an uploaded CV/resume file. Supports PDF,
DOCX, and plain TXT. Also does light structural extraction (a guess
at the candidate's name, and an email/phone check) purely from
formatting heuristics, used to personalize the report.
"""

import io
import re

import docx
import pdfplumber


def extract_text(uploaded_file) -> str:
    """
    Accepts a Streamlit UploadedFile (or any file-like object with
    a `.name` and readable bytes) and returns extracted plain text.
    """
    name = getattr(uploaded_file, "name", "").lower()
    raw_bytes = uploaded_file.read()

    if name.endswith(".pdf"):
        return _extract_pdf(raw_bytes)
    elif name.endswith(".docx"):
        return _extract_docx(raw_bytes)
    elif name.endswith(".txt"):
        return raw_bytes.decode("utf-8", errors="ignore")
    else:
        # Best-effort fallback
        try:
            return raw_bytes.decode("utf-8", errors="ignore")
        except Exception:
            return ""


def _extract_pdf(raw_bytes: bytes) -> str:
    text_chunks = []
    with pdfplumber.open(io.BytesIO(raw_bytes)) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text_chunks.append(page_text)
    return "\n".join(text_chunks)


def _extract_docx(raw_bytes: bytes) -> str:
    document = docx.Document(io.BytesIO(raw_bytes))
    paragraphs = [p.text for p in document.paragraphs]
    # Also pull text from tables (many resumes use table layouts)
    for table in document.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text:
                    paragraphs.append(cell.text)
    return "\n".join(paragraphs)


def guess_candidate_name(text: str, filename: str = "") -> str:
    """
    Heuristic name guesser: most resumes put the candidate's name on
    the first non-empty line, in Title Case, without digits/emails.
    Falls back to a cleaned-up version of the filename.
    """
    lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
    for line in lines[:6]:
        if "@" in line or re.search(r"\d{3,}", line):
            continue
        words = line.split()
        if 1 <= len(words) <= 4 and all(w[0:1].isupper() for w in words if w[0:1].isalpha()):
            if len(line) < 40:
                return line.title()

    # Fallback: use filename
    base = re.sub(r"\.(pdf|docx|txt)$", "", filename, flags=re.IGNORECASE)
    base = re.sub(r"[_\-]+", " ", base)
    base = re.sub(r"(cv|resume)", "", base, flags=re.IGNORECASE).strip()
    return base.title() if base else "Candidate"


def find_contact_info(text: str) -> dict:
    email_match = re.search(r"[\w.+-]+@[\w-]+\.[\w.-]+", text)
    phone_match = re.search(r"(\+?\d[\d\s\-()]{8,}\d)", text)
    return {
        "email": email_match.group(0) if email_match else None,
        "phone": phone_match.group(0).strip() if phone_match else None,
    }
